from docx import Document
from pathlib import Path


def main():
    docx_path = Path("Gilivindu Web Content.docx")
    out_path = Path("extracted_docx.txt")

    doc = Document(docx_path)
    lines: list[str] = []

    lines.append(f"FILE: {docx_path.name}")
    lines.append("")

    lines.append("=== PARAGRAPHS ===")
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style = ""
        try:
            style = p.style.name or ""
        except Exception:
            style = ""

        if style and style.lower() not in ("normal",):
            lines.append(f"[{style}] {text}")
        else:
            lines.append(text)

    if doc.tables:
        lines.append("")
        lines.append("=== TABLES ===")
        for ti, table in enumerate(doc.tables, start=1):
            lines.append(f"-- Table {ti} --")
            for row in table.rows:
                cells = [" ".join((c.text or "").split()) for c in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))

    out_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"Wrote {out_path.resolve()}")


if __name__ == "__main__":
    main()
